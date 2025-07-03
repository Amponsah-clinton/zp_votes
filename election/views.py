from django.shortcuts import render, redirect
from django.contrib import messages
from .forms import InitialsLoginForm
from .models import Learner, Position, Aspirant, Vote

from django.shortcuts import render, redirect
from django.contrib import messages
from .forms import InitialsLoginForm
from .models import Learner

from django.contrib import messages
from django.shortcuts import render, redirect
from .forms import InitialsLoginForm
from .models import Learner

from .models import Learner, Staff  # import Staff

def login_view(request):
    if request.method == 'POST':
        form = InitialsLoginForm(request.POST)
        if form.is_valid():
            initials = form.cleaned_data['initials'].strip().upper()
            password = form.cleaned_data['password']

            # Check if Learner exists
            try:
                learner = Learner.objects.get(initials=initials, password=password)
                if learner.has_voted:
                    messages.error(request, 'You have already voted.')
                    return redirect('login')
                request.session['voter_id'] = learner.id
                request.session['voter_type'] = 'learner'
                return redirect('vote')
            except Learner.DoesNotExist:
                pass

            # Check if Staff exists
            try:
                staff = Staff.objects.get(initials=initials, password=password)
                if staff.has_voted:
                    messages.error(request, 'You have already voted.')
                    return redirect('login')
                request.session['voter_id'] = staff.id
                request.session['voter_type'] = 'staff'
                return redirect('vote')
            except Staff.DoesNotExist:
                pass

            messages.error(request, 'Invalid initials or password.')
    else:
        form = InitialsLoginForm()

    return render(request, 'login.html', {'form': form})


from django.shortcuts import redirect, render
from django.contrib import messages
from .models import Learner, Staff, Vote, Position, Aspirant

def vote_view(request):
    voter_id = request.session.get('voter_id')
    voter_type = request.session.get('voter_type')

    if not voter_id or not voter_type:
        messages.error(request, 'Access Denied. You are either not logged in or have already voted.')
        return redirect('login')

    # Identify voter
    if voter_type == 'learner':
        try:
            voter = Learner.objects.get(id=voter_id)
        except Learner.DoesNotExist:
            messages.error(request, 'Invalid learner session.')
            return redirect('login')
    elif voter_type == 'staff':
        try:
            voter = Staff.objects.get(id=voter_id)
        except Staff.DoesNotExist:
            messages.error(request, 'Invalid staff session.')
            return redirect('login')
    else:
        messages.error(request, 'Unknown voter type.')
        return redirect('login')

    # Check voting status
    if voter.has_voted:
        messages.error(request, 'You have already voted.')
        return redirect('login')

    positions = Position.objects.all().order_by('id')

    if request.method == 'POST':
        for position in positions:
            aspirant_id = request.POST.get(str(position.id))
            if aspirant_id:
                aspirant = Aspirant.objects.get(id=aspirant_id)

                if voter_type == 'staff':
                    Vote.objects.create(
                        staff=voter,
                        position=position,
                        aspirant=aspirant
                    )
                else:
                    Vote.objects.create(
                        learner=voter,
                        position=position,
                        aspirant=aspirant
                    )

        voter.has_voted = True
        voter.save()

        return redirect('thank_you')

    position_aspirants = [
        (position, Aspirant.objects.filter(position=position)) for position in positions
    ]

    return render(request, 'vote.html', {
        'voter': voter,
        'positions': position_aspirants,
        'voter_type': voter_type,
    })




def stats_view(request):
    positions = Position.objects.all()

    data = []
    for position in positions:
        aspirants = Aspirant.objects.filter(position=position)
        position_votes = Vote.objects.filter(position=position).count()
        aspirant_data = []

        for asp in aspirants:
            votes = Vote.objects.filter(position=position, aspirant=asp).count()
            percentage = (votes / position_votes * 100) if position_votes > 0 else 0
            aspirant_data.append({
                'name': asp.name,
                'grade': asp.grade,
                'photo': asp.photo.url if asp.photo else '',
                'votes': votes,
                'percentage': round(percentage, 2),
            })

        aspirant_data.sort(key=lambda x: x['votes'], reverse=True)

        data.append({
            'position': position.name,
            'aspirants': aspirant_data,
            'total_votes': position_votes  # ✅ total for that position only
        })

    return render(request, 'stats.html', {'data': data})



from django.shortcuts import render
from .models import Position, Aspirant, Vote
from django.http import HttpResponse
import csv

def vote_report_view(request):
    positions = Position.objects.all()
    selected_position_id = request.GET.get('position')
    selected_aspirant_id = request.GET.get('aspirant')

    aspirants = Aspirant.objects.filter(position_id=selected_position_id) if selected_position_id else []
    votes = Vote.objects.none()

    if selected_aspirant_id:
        votes = Vote.objects.filter(aspirant_id=selected_aspirant_id).select_related('learner', 'aspirant', 'position')

    all_votes = Vote.objects.select_related('learner', 'aspirant', 'position').all()

    context = {
        'positions': positions,
        'aspirants': aspirants,
        'selected_position_id': selected_position_id,
        'selected_aspirant_id': selected_aspirant_id,
        'votes': votes,
        'all_votes': all_votes,
    }
    return render(request, 'vote_report.html', context)



from django.http import HttpResponse
from .models import Vote, Position, Learner, Staff
import csv

def export_votes_csv(request):
    positions = Position.objects.all()
    learners = Learner.objects.filter(has_voted=True)
    staff_members = Staff.objects.filter(has_voted=True)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="structured_votes.csv"'

    writer = csv.writer(response)

    # Header
    headers = ['Voter Name', 'Voter Type', 'Class/Role'] + [pos.name for pos in positions]
    writer.writerow(headers)

    # Write learners' votes
    for learner in learners:
        row = [learner.full_name, 'Learner', learner.grade]
        for pos in positions:
            vote = Vote.objects.filter(learner=learner, position=pos).first()
            row.append(vote.aspirant.name if vote else '-')
        writer.writerow(row)

    # Write staff votes
    for staff in staff_members:
        row = [staff.name, 'Staff', staff.role_or_subject]
        for pos in positions:
            vote = Vote.objects.filter(staff=staff, position=pos).first()
            row.append(vote.aspirant.name if vote else '-')
        writer.writerow(row)

    return response


from django.shortcuts import render, redirect
from .models import Learner, Staff

def thank_you(request):
    voter_id = request.session.get('voter_id')
    voter_type = request.session.get('voter_type')

    if not voter_id or not voter_type:
        return redirect('login')

    if voter_type == 'learner':
        voter = Learner.objects.get(id=voter_id)
    elif voter_type == 'staff':
        voter = Staff.objects.get(id=voter_id)
    else:
        return redirect('login')

    return render(request, 'thank_you.html', {
        'voter': voter,
        'voter_type': voter_type
    })

import csv
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Learner

def upload_learners_view(request):
    if request.method == 'POST' and request.FILES.get('csv_file'):
        file = request.FILES['csv_file']
        if not file.name.endswith('.csv'):
            messages.error(request, 'Please upload a .csv file.')
            return redirect('upload_learners')

        try:
            decoded_file = file.read().decode('utf-8').splitlines()
            reader = csv.DictReader(decoded_file)

            added = 0
            skipped = 0
            skipped_names = []

            for row in reader:
                full_name = row['full_name'].strip()
                grade = row['grade'].strip()

                # Generate initials (uppercase)
                initials = ''.join([n[0].upper() for n in full_name.split()])

                # Check if initials already exist in same grade
                if Learner.objects.filter(initials=initials, grade=grade).exists():
                    skipped += 1
                    skipped_names.append(full_name)
                    continue

                # Create learner
                Learner.objects.create(full_name=full_name, grade=grade, initials=initials)
                added += 1

            # Compose messages
            success_msg = f"Upload complete: {added} added."
            if skipped_names:
                skipped_msg = f"Skipped {skipped} due to duplicate initials in class: " + ", ".join(skipped_names)
                messages.warning(request, skipped_msg)
            messages.success(request, success_msg)

        except Exception as e:
            messages.error(request, f"Error processing file: {e}")

        return redirect('upload_learners')

    return render(request, 'upload_learners.html')



from docx import Document
from docx.shared import Inches
from django.http import HttpResponse
from django.conf import settings
from .models import Position, Aspirant, Vote, Learner

from PIL import Image, ExifTags
import os

def correct_image_orientation(image_path):
    """
    Auto-rotates the image based on EXIF orientation and saves a corrected copy.
    Returns path to the corrected image.
    """
    try:
        image = Image.open(image_path)

        # Look for EXIF orientation tag
        for orientation in ExifTags.TAGS.keys():
            if ExifTags.TAGS[orientation] == 'Orientation':
                break

        exif = image._getexif()
        if exif is not None:
            orientation_value = exif.get(orientation)
            if orientation_value == 3:
                image = image.rotate(180, expand=True)
            elif orientation_value == 6:
                image = image.rotate(270, expand=True)
            elif orientation_value == 8:
                image = image.rotate(90, expand=True)

        corrected_path = image_path + "_corrected.jpg"
        image.save(corrected_path)
        return corrected_path
    except Exception as e:
        print("Image correction failed:", e)
        return None

import os
import tempfile
from PIL import Image
from django.conf import settings
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
from .models import Position, Aspirant, Vote, Learner


def resize_and_convert_image(image_path, max_size=(300, 300)):
    """
    Resizes and converts image to JPEG RGB format for Word document use.
    """
    try:
        img = Image.open(image_path)
        img = img.convert('RGB')
        img.thumbnail(max_size)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
        img.save(temp_file.name, format='JPEG', quality=85)
        return temp_file.name
    except Exception as e:
        print("⚠️ Image resize error:", e)
        return None


import csv
import io
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Vote, Learner, Staff, Aspirant, Position
from django.core.exceptions import ValidationError

def upload_votes(request):
    if request.method == 'POST':
        csv_file = request.FILES.get('csv_file')
        if not csv_file.name.endswith('.csv'):
            messages.error(request, 'Please upload a valid CSV file.')
            return redirect('upload_votes')

        data_set = csv_file.read().decode('utf-8')
        io_string = io.StringIO(data_set)
        reader = csv.DictReader(io_string)

        success_count = 0

        for row in reader:
            voter_name = row.get('Voter Name', '').strip()
            voter_type = row.get('Voter Type', '').strip().lower()
            class_or_role = row.get('Class/Role', '').strip()

            learner = None
            staff = None

            # Identify learner or staff
            try:
                if voter_type == 'learner':
                    learner = Learner.objects.get(full_name__iexact=voter_name)
                elif voter_type == 'staff':
                    staff = Staff.objects.get(name__iexact=voter_name)
                else:
                    raise ValueError("Invalid Voter Type")
            except Exception as e:
                messages.warning(request, f"Could not find {voter_type} named '{voter_name}': {e}")
                continue

            # Process each position/aspirant vote
            for position_title, aspirant_name in row.items():
                if position_title in ['Voter Name', 'Voter Type', 'Class/Role'] or not aspirant_name.strip():
                    continue

                aspirant_name = aspirant_name.strip()
                position_title = position_title.strip()

                try:
                    position = Position.objects.get(name__iexact=position_title)
                    aspirant = Aspirant.objects.get(name__iexact=aspirant_name, position=position)

                    vote = Vote(
                        learner=learner,
                        staff=staff,
                        position=position,
                        aspirant=aspirant
                    )
                    vote.clean()
                    vote.save()
                    success_count += 1
                except Exception as e:
                    messages.warning(request, f"Error processing vote for '{aspirant_name}' as {position_title}: {e}")
                    continue

        messages.success(request, f"Uploaded {success_count} votes successfully.")
        return redirect('upload_votes')

    return render(request, 'upload_votes.html')




def generate_election_report(request):
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    document.add_heading('Zion Praise Educational Complex', 0)
    document.add_paragraph('🗳️ Official Election Report – Hillside Campus\n')

    static_logo_path = os.path.join(settings.BASE_DIR, 'static', 'images', 'logo.png')
    if os.path.exists(static_logo_path):
        resized_logo = resize_and_convert_image(static_logo_path, max_size=(400, 400))
        if resized_logo:
            document.add_picture(resized_logo, width=Inches(2.0))
            os.remove(resized_logo)

    document.add_paragraph("This document summarizes the votes, candidates, and winners from the student elections.")
    document.add_paragraph("Prepared and generated by the voting system.\n")
    document.add_page_break()

    learners_voted = Learner.objects.filter(has_voted=True).count()
    positions = Position.objects.all()

    for position in positions:
        document.add_heading(f"Position: {position.name}", level=1)

        aspirants = Aspirant.objects.filter(position=position)
        aspirants = sorted(aspirants, key=lambda a: Vote.objects.filter(position=position, aspirant=a).count(), reverse=True)

        for asp in aspirants:
            document.add_paragraph(f"{asp.name} (Class {asp.grade})", style='List Bullet')

            if asp.photo and hasattr(asp.photo, 'path') and os.path.exists(asp.photo.path):
                resized_image = resize_and_convert_image(asp.photo.path, max_size=(400, 400))
                if resized_image:
                    para = document.add_paragraph()
                    para.alignment = 1  # Center image
                    para.add_run().add_picture(resized_image, width=Inches(1.4))
                    os.remove(resized_image)
                else:
                    document.add_paragraph("📸 Image error.")
            else:
                document.add_paragraph("📸 No image available.")

            votes = Vote.objects.filter(position=position, aspirant=asp).count()
            percent = (votes / learners_voted * 100) if learners_voted > 0 else 0
            document.add_paragraph(f"Votes: {votes} ({round(percent, 2)}%)\n")

        document.add_page_break()

    # === Winner Summary Table ===
    document.add_heading('🏆 Winners Summary', level=1)

    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'

    headers = ['Position', 'Role', 'Photo', 'Name', 'Class', 'Votes', 'Percentage']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    single_role_positions = [
        "Assistant Head Prefect - Girls",
        "Assistant Head Prefect - Boys",
        "Head Prefect - Girls",
        "Head Prefect - Boys"
    ]

    def add_winner_row(role, asp, position_name):
        row = table.add_row().cells
        row[0].text = position_name
        row[1].text = role

        # Add photo if available
        if asp.photo and hasattr(asp.photo, 'path') and os.path.exists(asp.photo.path):
            resized = resize_and_convert_image(asp.photo.path, max_size=(300, 300))
            if resized:
                try:
                    run = row[2].paragraphs[0].add_run()
                    run.add_picture(resized, width=Inches(1.2))
                    os.remove(resized)
                except:
                    row[2].text = 'Image error'
            else:
                row[2].text = 'Image error'
        else:
            row[2].text = 'No image'

        row[3].text = asp.name
        row[4].text = asp.grade
        votes = Vote.objects.filter(position__name=position_name, aspirant=asp).count()
        percent = (votes / learners_voted * 100) if learners_voted > 0 else 0
        row[5].text = str(votes)
        row[6].text = f"{round(percent, 2)}%"

    for position in positions:
        aspirants = Aspirant.objects.filter(position=position)
        aspirants = sorted(aspirants, key=lambda a: Vote.objects.filter(position=position, aspirant=a).count(), reverse=True)

        if not aspirants:
            continue

        if position.name in single_role_positions:
            add_winner_row("Prefect", aspirants[0], position.name)
        else:
            add_winner_row("Main Prefect", aspirants[0], position.name)
            if len(aspirants) > 1:
                add_winner_row("Assistant Prefect", aspirants[1], position.name)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="ZPEC_Election_Report.docx"'
    document.save(response)
    return response



from django.shortcuts import render

def report_page_view(request):
    return render(request, 'report_page.html')

def adminpage(request):
    return render(request, 'adminpage.html')



from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from .models import Learner
from collections import defaultdict

def download_learners_doc(request):
    # Group learners by grade
    learners_by_grade = defaultdict(list)
    for learner in Learner.objects.all().order_by('grade', 'full_name'):
        learners_by_grade[learner.grade].append(learner)

    # Create the Word document
    doc = Document()
    doc.add_heading('Zion Praise Educational Complex – Learner List', 0)

    for grade, learners in learners_by_grade.items():
        doc.add_heading(f'Class {grade}', level=1)

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No.'
        hdr_cells[1].text = 'Full Name'
        hdr_cells[2].text = 'Initials'

        for idx, learner in enumerate(learners, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = learner.full_name
            row_cells[2].text = learner.initials

        doc.add_paragraph()  

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=learners_by_class.docx'
    doc.save(response)
    return response




from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import Learner

def download_learners_word(request):
    from math import ceil

    grades = ['7A', '7B', '7C', '8A', '8B']
    doc = Document()

    # Set A4 margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    for idx, grade in enumerate(grades):
        learners = list(Learner.objects.filter(grade=grade).order_by('full_name'))

        doc.add_heading(f'ZPEC Voting Slips - Class {grade}', level=1)

        # Create a table with 3 columns
        num_rows = ceil(len(learners) / 3)
        table = doc.add_table(rows=num_rows, cols=3)
        table.autofit = True

        col = 0
        row = 0
        for learner in learners:
            cell = table.cell(row, col)
            para = cell.paragraphs[0]

            run = para.add_run(f"{learner.full_name}\n")
            run.bold = True
            run.font.size = Pt(10)

            para.add_run(f"Initials: {learner.initials}\n").font.size = Pt(9)
            para.add_run(f"Password: {learner.password}\n").font.size = Pt(9)
            para.add_run("-" * 25 + "\n").font.size = Pt(8)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            col += 1
            if col >= 3:
                col = 0
                row += 1

        # Add a page break unless it's the last one
        if idx < len(grades) - 1:
            doc.add_page_break()

    # Create and return Word doc
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=zpec_learner_slips.docx'
    doc.save(response)
    return response



from django.http import HttpResponse, HttpResponseServerError
from docx import Document
from docx.shared import Pt
from django.db.models import Count
from .models import Position, Aspirant

def download_results_summary(request):
    try:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        doc.add_heading('🏆 ZPEC Voting Summary Report', 0)

        positions = Position.objects.all()

        for position in positions:
            doc.add_heading(f'Position: {position.name}', level=2)

            aspirants = Aspirant.objects.filter(position=position) \
                .annotate(vote_count=Count('votes')) \
                .order_by('-vote_count')  # Highest votes on top

            total_votes = sum(a.vote_count for a in aspirants) or 1  # avoid division by zero

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Aspirant'
            hdr_cells[1].text = 'Votes'
            hdr_cells[2].text = 'Percentage'

            max_votes = aspirants[0].vote_count if aspirants else 0

            for aspirant in aspirants:
                percent = (aspirant.vote_count / total_votes) * 100
                row_cells = table.add_row().cells
                winner_icon = ' 🏆' if aspirant.vote_count == max_votes and max_votes > 0 else ''
                row_cells[0].text = aspirant.name + winner_icon
                row_cells[1].text = f"{aspirant.vote_count}"
                row_cells[2].text = f"{percent:.2f}%"

            # Add total row at bottom
            total_row = table.add_row().cells
            total_row[0].text = "Total Votes"
            total_row[1].text = f"{total_votes} "
            total_row[2].text = "100.00%"

            doc.add_paragraph('\n')

        # Word response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=zpec_voting_results.docx'
        doc.save(response)
        return response

    except Exception as e:
        print("🔴 Error in download_results_summary:", e)
        return HttpResponseServerError("Something went wrong: " + str(e))




from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from math import ceil
from .models import Staff

def download_staff_credentials_word(request):
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    staff_members = list(Staff.objects.all().order_by('name'))

    doc.add_heading('ZPEC Voting Slips - Staff Credentials', level=1)

    num_rows = ceil(len(staff_members) / 3)
    table = doc.add_table(rows=num_rows, cols=3)
    table.autofit = True

    col = 0
    row = 0
    for staff in staff_members:
        cell = table.cell(row, col)
        para = cell.paragraphs[0]

        run = para.add_run(f"{staff.title} {staff.name}\n")
        run.bold = True
        run.font.size = Pt(10)

        para.add_run(f"Initials: {staff.initials}\n").font.size = Pt(9)
        para.add_run(f"Password: {staff.password}\n").font.size = Pt(9)
        para.add_run("-" * 25 + "\n").font.size = Pt(8)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        col += 1
        if col >= 3:
            col = 0
            row += 1

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=zpec_staff_credentials.docx'
    doc.save(response)
    return response
